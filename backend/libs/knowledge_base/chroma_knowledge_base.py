"""
PeerPortal 基于 Chroma 云数据库的知识库系统
支持文档摄取、向量存储、语义搜索等功能
"""
import os
import uuid
import asyncio
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
from datetime import datetime
import hashlib

import chromadb
from chromadb.config import Settings
import PyPDF2
import docx
from bs4 import BeautifulSoup
import markdown

# 如果没有安装某些库，提供降级方案
try:
    import openai
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False

try:
    from sentence_transformers import SentenceTransformer
    HAS_SENTENCE_TRANSFORMERS = True
except ImportError:
    HAS_SENTENCE_TRANSFORMERS = False


class Document:
    """文档对象"""
    def __init__(self, content: str, metadata: Dict[str, Any] = None):
        self.id = str(uuid.uuid4())
        self.content = content
        self.metadata = metadata or {}
        self.created_at = datetime.now().isoformat()


class DocumentChunk:
    """文档块对象"""
    def __init__(self, content: str, metadata: Dict[str, Any] = None, chunk_id: str = None):
        self.id = chunk_id or str(uuid.uuid4())
        self.content = content
        self.metadata = metadata or {}
        self.created_at = datetime.now().isoformat()


class DocumentLoader:
    """文档加载器"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    async def load_file(self, file_path: str, metadata: Dict[str, Any] = None) -> Document:
        """加载文件"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 基础元数据
        file_metadata = {
            "filename": file_path.name,
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size,
            "file_extension": file_path.suffix.lower(),
            "created_at": datetime.now().isoformat(),
            **(metadata or {})
        }
        
        # 根据文件类型加载内容
        if file_path.suffix.lower() == '.pdf':
            content = await self._load_pdf(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            content = await self._load_docx(file_path)
        elif file_path.suffix.lower() == '.txt':
            content = await self._load_txt(file_path)
        elif file_path.suffix.lower() == '.md':
            content = await self._load_markdown(file_path)
        elif file_path.suffix.lower() in ['.html', '.htm']:
            content = await self._load_html(file_path)
        else:
            # 尝试作为文本文件读取
            try:
                content = await self._load_txt(file_path)
            except Exception as e:
                raise ValueError(f"不支持的文件类型: {file_path.suffix}，错误: {e}")
        
        return Document(content=content, metadata=file_metadata)
    
    async def _load_pdf(self, file_path: Path) -> str:
        """加载PDF文件"""
        try:
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content.append(page.extract_text())
            return '\\n\\n'.join(content)
        except Exception as e:
            self.logger.warning(f"PDF加载失败，尝试文本模式: {e}")
            # 降级到文本模式
            return await self._load_txt(file_path)
    
    async def _load_docx(self, file_path: Path) -> str:
        """加载DOCX文件"""
        try:
            doc = docx.Document(file_path)
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            return '\\n\\n'.join(content)
        except Exception as e:
            self.logger.warning(f"DOCX加载失败，尝试文本模式: {e}")
            return await self._load_txt(file_path)
    
    async def _load_txt(self, file_path: Path) -> str:
        """加载文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"无法解码文件: {file_path}")
    
    async def _load_markdown(self, file_path: Path) -> str:
        """加载Markdown文件"""
        try:
            content = await self._load_txt(file_path)
            # 转换为HTML再提取纯文本
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
        except Exception as e:
            self.logger.warning(f"Markdown解析失败，返回原文本: {e}")
            return await self._load_txt(file_path)
    
    async def _load_html(self, file_path: Path) -> str:
        """加载HTML文件"""
        try:
            html_content = await self._load_txt(file_path)
            soup = BeautifulSoup(html_content, 'html.parser')
            return soup.get_text()
        except Exception as e:
            self.logger.warning(f"HTML解析失败，返回原文本: {e}")
            return await self._load_txt(file_path)


class TextSplitter:
    """文本分割器"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.logger = logging.getLogger(__name__)
    
    def split_document(self, document: Document) -> List[DocumentChunk]:
        """分割文档为块"""
        chunks = []
        text = document.content
        
        # 首先按段落分割
        paragraphs = [p.strip() for p in text.split('\\n\\n') if p.strip()]
        
        current_chunk = ""
        chunk_index = 0
        
        for paragraph in paragraphs:
            # 如果当前块加上新段落不超过限制，则添加
            if len(current_chunk) + len(paragraph) + 2 <= self.chunk_size:
                if current_chunk:
                    current_chunk += "\\n\\n"
                current_chunk += paragraph
            else:
                # 保存当前块
                if current_chunk:
                    chunk = self._create_chunk(
                        current_chunk, 
                        document.metadata, 
                        chunk_index
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                
                # 如果段落本身太长，需要进一步分割
                if len(paragraph) > self.chunk_size:
                    sub_chunks = self._split_long_text(paragraph)
                    for sub_chunk in sub_chunks:
                        chunk = self._create_chunk(
                            sub_chunk, 
                            document.metadata, 
                            chunk_index
                        )
                        chunks.append(chunk)
                        chunk_index += 1
                    current_chunk = ""
                else:
                    # 开始新的块
                    current_chunk = paragraph
        
        # 处理最后一个块
        if current_chunk:
            chunk = self._create_chunk(current_chunk, document.metadata, chunk_index)
            chunks.append(chunk)
        
        self.logger.info(f"文档分割完成，共生成 {len(chunks)} 个块")
        return chunks
    
    def _create_chunk(self, content: str, base_metadata: Dict, chunk_index: int) -> DocumentChunk:
        """创建文档块"""
        chunk_metadata = {
            **base_metadata,
            "chunk_index": chunk_index,
            "chunk_size": len(content),
            "chunk_id": f"{base_metadata.get('filename', 'unknown')}_{chunk_index}"
        }
        
        return DocumentChunk(
            content=content,
            metadata=chunk_metadata,
            chunk_id=chunk_metadata["chunk_id"]
        )
    
    def _split_long_text(self, text: str) -> List[str]:
        """分割过长的文本"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # 尝试在句号或换行符处分割
            split_pos = end
            for i in range(end - 100, end):
                if i > start and text[i] in '.。\\n':
                    split_pos = i + 1
                    break
            
            chunks.append(text[start:split_pos])
            start = split_pos - self.chunk_overlap
        
        return chunks


class EmbeddingProvider:
    """嵌入向量提供器"""
    
    def __init__(self, provider: str = "sentence_transformers", model_name: str = None):
        self.provider = provider
        self.model_name = model_name
        self.logger = logging.getLogger(__name__)
        self._model = None
        
        self._initialize_model()
    
    def _initialize_model(self):
        """初始化嵌入模型"""
        if self.provider == "openai" and HAS_OPENAI:
            self.model_name = self.model_name or "text-embedding-ada-002"
            # OpenAI 不需要预加载模型
        elif self.provider == "sentence_transformers" and HAS_SENTENCE_TRANSFORMERS:
            self.model_name = self.model_name or "paraphrase-multilingual-MiniLM-L12-v2"
            try:
                self._model = SentenceTransformer(self.model_name)
                self.logger.info(f"已加载 Sentence Transformers 模型: {self.model_name}")
            except Exception as e:
                self.logger.warning(f"Sentence Transformers 模型加载失败: {e}")
                self._model = None
        else:
            self.logger.warning("没有可用的嵌入模型，将使用模拟向量")
    
    async def embed_texts(self, texts: List[str]) -> List[List[float]]:
        """生成文本嵌入向量"""
        if self.provider == "openai" and HAS_OPENAI and os.getenv("OPENAI_API_KEY"):
            return await self._embed_with_openai(texts)
        elif self.provider == "sentence_transformers" and self._model:
            return await self._embed_with_sentence_transformers(texts)
        else:
            return await self._embed_with_mock(texts)
    
    async def _embed_with_openai(self, texts: List[str]) -> List[List[float]]:
        """使用 OpenAI 生成嵌入"""
        try:
            client = openai.OpenAI()
            response = client.embeddings.create(
                model=self.model_name,
                input=texts
            )
            return [item.embedding for item in response.data]
        except Exception as e:
            self.logger.warning(f"OpenAI 嵌入生成失败: {e}，使用模拟向量")
            return await self._embed_with_mock(texts)
    
    async def _embed_with_sentence_transformers(self, texts: List[str]) -> List[List[float]]:
        """使用 Sentence Transformers 生成嵌入"""
        try:
            embeddings = self._model.encode(texts)
            return embeddings.tolist()
        except Exception as e:
            self.logger.warning(f"Sentence Transformers 嵌入生成失败: {e}，使用模拟向量")
            return await self._embed_with_mock(texts)
    
    async def _embed_with_mock(self, texts: List[str]) -> List[List[float]]:
        """生成模拟嵌入向量"""
        embeddings = []
        for text in texts:
            # 基于文本内容生成确定性的模拟向量
            hash_obj = hashlib.md5(text.encode())
            hash_bytes = hash_obj.digest()
            
            # 将哈希值转换为384维向量 (模拟 sentence-transformers 的维度)
            vector = []
            for i in range(0, len(hash_bytes)):
                # 每个字节生成24个浮点数
                for j in range(24):
                    val = (hash_bytes[i] + j) / 255.0 - 0.5  # 归一化到[-0.5, 0.5]
                    vector.append(val)
            
            # 确保向量长度为384
            while len(vector) < 384:
                vector.extend(vector[:384-len(vector)])
            vector = vector[:384]
            
            embeddings.append(vector)
        
        return embeddings


class ChromaKnowledgeBase:
    """基于 Chroma 云数据库的知识库"""
    
    def __init__(
        self, 
        api_key: str,
        tenant: str,
        database: str,
        collection_name: str = "documents",
        embedding_provider: str = "sentence_transformers"
    ):
        self.api_key = api_key
        self.tenant = tenant
        self.database = database
        self.collection_name = collection_name
        
        # 初始化组件
        self.document_loader = DocumentLoader()
        self.text_splitter = TextSplitter()
        self.embedding_provider = EmbeddingProvider(embedding_provider)
        self.logger = logging.getLogger(__name__)
        
        # 初始化 Chroma 客户端
        self._initialize_client()
    
    def _initialize_client(self):
        """初始化 Chroma 客户端"""
        try:
            self.client = chromadb.CloudClient(
                api_key=self.api_key,
                tenant=self.tenant,
                database=self.database
            )
            
            # 获取或创建集合
            try:
                self.collection = self.client.get_collection(self.collection_name)
                self.logger.info(f"已连接到现有集合: {self.collection_name}")
            except Exception:
                # 集合不存在，创建新的
                self.collection = self.client.create_collection(
                    name=self.collection_name,
                    metadata={"description": "PeerPortal 知识库文档集合"}
                )
                self.logger.info(f"已创建新集合: {self.collection_name}")
                
        except Exception as e:
            self.logger.error(f"Chroma 客户端初始化失败: {e}")
            raise
    
    async def add_document(
        self, 
        file_path: str, 
        metadata: Dict[str, Any] = None,
        tenant_id: str = None
    ) -> Dict[str, Any]:
        """添加文档到知识库"""
        try:
            self.logger.info(f"开始处理文档: {file_path}")
            
            # 1. 加载文档
            document = await self.document_loader.load_file(file_path, metadata)
            
            # 添加租户信息
            if tenant_id:
                document.metadata["tenant_id"] = tenant_id
            
            # 2. 分割文档
            chunks = self.text_splitter.split_document(document)
            
            if not chunks:
                return {
                    "success": False,
                    "error": "文档分割后没有生成有效的块",
                    "document_id": None,
                    "chunks_count": 0
                }
            
            # 3. 生成嵌入向量
            chunk_texts = [chunk.content for chunk in chunks]
            embeddings = await self.embedding_provider.embed_texts(chunk_texts)
            
            # 4. 准备数据
            document_id = str(uuid.uuid4())
            ids = [f"{document_id}_{i}" for i in range(len(chunks))]
            documents = chunk_texts
            metadatas = []
            
            for i, chunk in enumerate(chunks):
                chunk_metadata = {
                    **chunk.metadata,
                    "document_id": document_id,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
                metadatas.append(chunk_metadata)
            
            # 5. 存储到 Chroma
            self.collection.add(
                ids=ids,
                documents=documents,
                embeddings=embeddings,
                metadatas=metadatas
            )
            
            self.logger.info(f"文档添加成功: {file_path}, 生成 {len(chunks)} 个块")
            
            return {
                "success": True,
                "document_id": document_id,
                "filename": document.metadata.get("filename"),
                "chunks_count": len(chunks),
                "file_size": document.metadata.get("file_size"),
                "embedding_model": self.embedding_provider.model_name
            }
            
        except Exception as e:
            self.logger.error(f"文档添加失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "document_id": None,
                "chunks_count": 0
            }
    
    async def search(
        self, 
        query: str, 
        top_k: int = 5,
        tenant_id: str = None,
        filter_metadata: Dict[str, Any] = None
    ) -> List[Dict[str, Any]]:
        """搜索知识库"""
        try:
            self.logger.info(f"搜索查询: {query}")
            
            # 1. 生成查询嵌入
            query_embeddings = await self.embedding_provider.embed_texts([query])
            query_embedding = query_embeddings[0]
            
            # 2. 构建过滤条件
            where_filter = {}
            if tenant_id:
                where_filter["tenant_id"] = tenant_id
            if filter_metadata:
                where_filter.update(filter_metadata)
            
            # 3. 执行搜索
            search_kwargs = {
                "query_embeddings": [query_embedding],
                "n_results": top_k,
                "include": ["documents", "metadatas", "distances"]
            }
            
            if where_filter:
                search_kwargs["where"] = where_filter
            
            results = self.collection.query(**search_kwargs)
            
            # 4. 格式化结果
            search_results = []
            if results["documents"] and results["documents"][0]:
                for i in range(len(results["documents"][0])):
                    result = {
                        "content": results["documents"][0][i],
                        "metadata": results["metadatas"][0][i],
                        "distance": results["distances"][0][i],
                        "score": 1 - results["distances"][0][i]  # 转换为相似度分数
                    }
                    search_results.append(result)
            
            self.logger.info(f"搜索完成，找到 {len(search_results)} 个结果")
            return search_results
            
        except Exception as e:
            self.logger.error(f"搜索失败: {e}")
            return []
    
    async def delete_document(self, document_id: str) -> bool:
        """删除文档"""
        try:
            # 查找所有属于该文档的块
            results = self.collection.query(
                query_embeddings=[[0] * 384],  # 虚拟查询
                n_results=1000,  # 获取更多结果
                where={"document_id": document_id},
                include=["metadatas"]
            )
            
            if results["ids"] and results["ids"][0]:
                chunk_ids = results["ids"][0]
                self.collection.delete(ids=chunk_ids)
                self.logger.info(f"已删除文档 {document_id}，共 {len(chunk_ids)} 个块")
                return True
            else:
                self.logger.warning(f"未找到文档: {document_id}")
                return False
                
        except Exception as e:
            self.logger.error(f"删除文档失败: {e}")
            return False
    
    def get_collection_info(self) -> Dict[str, Any]:
        """获取集合信息"""
        try:
            count = self.collection.count()
            return {
                "collection_name": self.collection_name,
                "document_count": count,
                "database": self.database,
                "tenant": self.tenant
            }
        except Exception as e:
            self.logger.error(f"获取集合信息失败: {e}")
            return {}
    
    async def list_documents(self, tenant_id: str = None, limit: int = 100) -> List[Dict[str, Any]]:
        """列出文档"""
        try:
            where_filter = {}
            if tenant_id:
                where_filter["tenant_id"] = tenant_id
            
            # 只获取每个文档的第一个块来代表文档
            where_filter["chunk_index"] = 0
            
            search_kwargs = {
                "query_embeddings": [[0] * 384],  # 虚拟查询
                "n_results": limit,
                "include": ["metadatas"],
                "where": where_filter
            }
            
            results = self.collection.query(**search_kwargs)
            
            documents = []
            if results["metadatas"] and results["metadatas"][0]:
                for metadata in results["metadatas"][0]:
                    doc_info = {
                        "document_id": metadata.get("document_id"),
                        "filename": metadata.get("filename"),
                        "file_size": metadata.get("file_size"),
                        "total_chunks": metadata.get("total_chunks"),
                        "created_at": metadata.get("created_at"),
                        "tenant_id": metadata.get("tenant_id")
                    }
                    documents.append(doc_info)
            
            return documents
            
        except Exception as e:
            self.logger.error(f"列出文档失败: {e}")
            return []


# 全局知识库实例
knowledge_base = None


def initialize_knowledge_base(
    api_key: str,
    tenant: str, 
    database: str,
    collection_name: str = "documents",
    embedding_provider: str = "sentence_transformers"
) -> ChromaKnowledgeBase:
    """初始化知识库"""
    global knowledge_base
    knowledge_base = ChromaKnowledgeBase(
        api_key=api_key,
        tenant=tenant,
        database=database,
        collection_name=collection_name,
        embedding_provider=embedding_provider
    )
    return knowledge_base


def get_knowledge_base() -> ChromaKnowledgeBase:
    """获取知识库实例"""
    if knowledge_base is None:
        raise RuntimeError("知识库未初始化，请先调用 initialize_knowledge_base()")
    return knowledge_base


# 留学项目搜索功能
def get_study_program_search():
    """获取留学项目搜索实例"""
    from .study_program_search import get_study_program_search
    return get_study_program_search()
